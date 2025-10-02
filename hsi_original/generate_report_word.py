#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import numpy as np
import tensorflow as tf
import matplotlib
matplotlib.use('Agg')  # headless
import matplotlib.pyplot as plt
from sklearn.metrics import (
    confusion_matrix, accuracy_score, precision_score, recall_score,
    f1_score, roc_auc_score, roc_curve, precision_recall_curve
)
from docx import Document
from docx.shared import Inches
import pandas as pd

# =============================
# CONFIG
# =============================
DATA_PATH = "./hsi_cnn"
MODEL_NAMES = {
    "Full": "model_cnn1d_full.keras",
    "Selected": "model_cnn1d_selected.keras"
}
BAND_SELECTION_FILE = os.path.join(DATA_PATH, "selected_indices.npy")

# -----------------------------
# Utilities
# -----------------------------
def load_data(selected: bool = False):
    """Load X and y; add channel if needed; fallback y_selected -> y."""
    X_file = "X_selected.npy" if selected else "X.npy"
    y_file = "y_selected.npy" if selected else "y.npy"

    X_path = os.path.join(DATA_PATH, X_file)
    y_path = os.path.join(DATA_PATH, y_file)

    if not os.path.exists(X_path):
        raise FileNotFoundError(f"X file not found: {X_path}")

    # Fallback: if y_selected.npy doesn't exist, use y.npy
    if not os.path.exists(y_path):
        y_path_alt = os.path.join(DATA_PATH, "y.npy")
        if os.path.exists(y_path_alt):
            y_path = y_path_alt
        else:
            raise FileNotFoundError(f"y file not found: {y_path} (or fallback {y_path_alt})")

    X = np.load(X_path)
    y = np.load(y_path)

    if X.ndim == 2:
        X = X[..., np.newaxis]  # ensure 3D for 1D-CNN

    return X, y

def save_confusion_matrix(cm: np.ndarray, classes, title: str, out_path: str):
    """Confusion matrix with cell grid and English labels."""
    plt.figure(figsize=(6, 5))
    plt.imshow(cm, interpolation='nearest', cmap=plt.cm.Blues)
    plt.title(title)
    cbar = plt.colorbar()
    cbar.ax.set_ylabel('Count', rotation=90)

    tick_marks = np.arange(len(classes))
    plt.xticks(tick_marks, classes, rotation=45, ha='right')
    plt.yticks(tick_marks, classes)

    thresh = cm.max() / 2.0 if cm.size else 0.5
    for i, j in np.ndindex(cm.shape):
        plt.text(j, i, int(cm[i, j]),
                 ha="center", va="center",
                 color="white" if cm[i, j] > thresh else "black")

    ax = plt.gca()
    # cell borders
    ax.set_xticks(np.arange(-0.5, len(classes), 1), minor=True)
    ax.set_yticks(np.arange(-0.5, len(classes), 1), minor=True)
    ax.grid(which='minor', color='white', linestyle='-', linewidth=0.8)
    ax.tick_params(which='minor', bottom=False, left=False)

    # general dashed grid
    plt.grid(True, linestyle='--', linewidth=0.5, alpha=0.4)

    plt.ylabel('True Label')
    plt.xlabel('Predicted Label')
    plt.tight_layout()
    plt.savefig(out_path, dpi=150)
    plt.close()

def save_roc_curve(y_true, y_score, title: str, out_path: str):
    fpr, tpr, _ = roc_curve(y_true, y_score)
    roc_auc = roc_auc_score(y_true, y_score)
    plt.figure(figsize=(6, 5))
    plt.plot(fpr, tpr, label=f'ROC curve (AUC = {roc_auc:.2f})')
    plt.plot([0, 1], [0, 1], 'k--', label='Chance')
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title(title)
    plt.legend(loc='lower right')
    plt.grid(True, linestyle='--', linewidth=0.5)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150)
    plt.close()
    return roc_auc

def save_prc_curve(y_true, y_score, title: str, out_path: str):
    precision, recall, _ = precision_recall_curve(y_true, y_score)
    plt.figure(figsize=(6, 5))
    plt.plot(recall, precision, label='Precision-Recall Curve')
    plt.xlabel('Recall')
    plt.ylabel('Precision')
    plt.title(title)
    plt.legend()
    plt.grid(True, linestyle='--', linewidth=0.5)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150)
    plt.close()

def save_error_spectra(X, y, y_pred, title: str, out_path: str):
    """Mean spectra for FP and FN; English labels + grid."""
    # X shape: (N, Bands, 1) or (N, Bands)
    X2d = X.squeeze()  # (N, Bands)
    if X2d.ndim != 2:
        return None  # safety

    mask_fp = (y == 0) & (y_pred == 1)
    mask_fn = (y == 1) & (y_pred == 0)

    fp_spectra = X2d[mask_fp].mean(axis=0) if np.any(mask_fp) else None
    fn_spectra = X2d[mask_fn].mean(axis=0) if np.any(mask_fn) else None

    plt.figure(figsize=(7, 4))
    any_line = False
    if fp_spectra is not None:
        plt.plot(fp_spectra, label='False Positives (mean)')
        any_line = True
    if fn_spectra is not None:
        plt.plot(fn_spectra, label='False Negatives (mean)')
        any_line = True

    if not any_line:
        plt.close()
        return None

    plt.title(title)
    plt.xlabel('Band Index')
    plt.ylabel('Amplitude (a.u.)')
    plt.legend()
    plt.grid(True, linestyle='--', linewidth=0.5)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150)
    plt.close()
    return out_path

def evaluate_model(name: str, model, X, y, band_indices=None):
    print(f"\n🚀 Processing model: {name}")

    y_pred_proba = model.predict(X, verbose=0).ravel()
    y_pred = (y_pred_proba > 0.5).astype(int)

    # Metrics
    acc = accuracy_score(y, y_pred)
    prec = precision_score(y, y_pred, zero_division=0)
    rec = recall_score(y, y_pred, zero_division=0)
    f1 = f1_score(y, y_pred, zero_division=0)
    roc_auc = roc_auc_score(y, y_pred_proba)

    print(f"Metrics — Acc: {acc:.4f} | Prec: {prec:.4f} | Rec: {rec:.4f} | F1: {f1:.4f} | ROC AUC: {roc_auc:.4f}")

    # Confusion Matrix (counts)
    cm = confusion_matrix(y, y_pred)
    classes = ["Others", "ATCC"]
    cm_path = f"confusion_matrix_{name}.png"
    save_confusion_matrix(cm, classes, f"Confusion Matrix — {name}", cm_path)

    # ROC
    roc_path = f"roc_{name}.png"
    _auc = save_roc_curve(y, y_pred_proba, f"ROC Curve — {name}", roc_path)  # reuses ROC AUC

    # Precision-Recall
    prc_path = f"prc_{name}.png"
    save_prc_curve(y, y_pred_proba, f"Precision–Recall Curve — {name}", prc_path)

    # Error spectra
    error_spectra_path = f"error_spectra_{name}.png"
    error_plot_path = save_error_spectra(X, y, y_pred, f"Error Spectra — {name}", error_spectra_path)

    # CSV exports
    pd.DataFrame(cm, columns=["Pred_Neg", "Pred_Pos"], index=["True_Neg", "True_Pos"]).to_csv(f"confusion_matrix_{name}.csv")
    pd.DataFrame({
        "Accuracy": [acc],
        "Precision": [prec],
        "Recall": [rec],
        "F1_Score": [f1],
        "ROC_AUC": [roc_auc]
    }).to_csv(f"metrics_{name}.csv", index=False)

    return {
        "name": name,
        "metrics": (acc, prec, rec, f1, roc_auc),
        "cm_path": cm_path,
        "roc_path": roc_path,
        "prc_path": prc_path,
        "error_spectra_path": error_plot_path,
        "band_indices": band_indices
    }

# =============================
# MAIN
# =============================
document = Document()
document.add_heading('CNN 1D Model Performance Report', 0)

# load band indices if available
band_indices = np.load(BAND_SELECTION_FILE) if os.path.exists(BAND_SELECTION_FILE) else None

results = []
for key, model_name in MODEL_NAMES.items():
    selected = (key == "Selected")
    X, y = load_data(selected)
    model = tf.keras.models.load_model(os.path.join(DATA_PATH, model_name))
    res = evaluate_model(key, model, X, y, band_indices)
    results.append(res)

# Build Word document
for res in results:
    document.add_heading(f'Model — {res["name"]}', level=1)

    acc, prec, rec, f1, roc_auc = res["metrics"]
    document.add_paragraph(
        f'Accuracy: {acc:.4f}\n'
        f'Precision: {prec:.4f}\n'
        f'Recall: {rec:.4f}\n'
        f'F1-Score: {f1:.4f}\n'
        f'ROC AUC: {roc_auc:.4f}\n'
    )

    # Figures
    document.add_heading('Figures', level=2)
    if res["cm_path"] and os.path.exists(res["cm_path"]):
        document.add_paragraph('Confusion Matrix')
        document.add_picture(res["cm_path"], width=Inches(4.5))
    if res["roc_path"] and os.path.exists(res["roc_path"]):
        document.add_paragraph('ROC Curve')
        document.add_picture(res["roc_path"], width=Inches(4.5))
    if res["prc_path"] and os.path.exists(res["prc_path"]):
        document.add_paragraph('Precision–Recall Curve')
        document.add_picture(res["prc_path"], width=Inches(4.5))
    if res["error_spectra_path"] and os.path.exists(res["error_spectra_path"]):
        document.add_paragraph('Error Spectra (mean)')
        document.add_picture(res["error_spectra_path"], width=Inches(4.5))

    # Selected bands (if available)
    if res["band_indices"] is not None and res["name"] == "Selected":
        document.add_heading('Selected Bands (indices)', level=2)
        # display first 30 per line for readability
        idx = res["band_indices"].astype(int).tolist()
        chunks = [idx[i:i+30] for i in range(0, len(idx), 30)]
        for ch in chunks:
            document.add_paragraph(' '.join(map(str, ch)))

document.save('CNN1D_Performance_Report.docx')
print("✅ Report generated: CNN1D_Performance_Report.docx")
